import os
import PyPDF2
from docx import Document
import re
from typing import Optional


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text content from various file formats (txt, pdf, docx)
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_ext == '.txt':
            return extract_text_from_txt(file_path)
        elif file_ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            return extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    except Exception as e:
        raise Exception(f"Failed to extract text from {file_path}: {str(e)}")

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from .txt files"""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
                return clean_text(text)
        except UnicodeDecodeError:
            continue
    
    raise Exception("Could not decode text file with any supported encoding")

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF files"""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            if num_pages == 0:
                raise Exception("PDF has no pages")
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if not text.strip():
            raise Exception("No text could be extracted from PDF")
        
        return clean_text(text)
    
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files"""
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            raise Exception("No text could be extracted from DOCX")
        
        return clean_text(text)
    
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")

def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text
    """
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove excessive newlines
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Remove special characters and control characters
    text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    
    # Strip leading/trailing whitespace
    text = text.strip()
    
    return text

def validate_file_size(file_path: str, max_size_mb: int = 10) -> bool:
    """
    Validate file size is within limits
    """
    max_size_bytes = max_size_mb * 1024 * 1024
    file_size = os.path.getsize(file_path)
    return file_size <= max_size_bytes

def validate_file_type(file_path: str) -> bool:
    """
    Validate file type is supported
    """
    allowed_extensions = ['.txt', '.pdf', '.docx']
    file_ext = os.path.splitext(file_path)[1].lower()
    return file_ext in allowed_extensions

def get_file_info(file_path: str) -> dict:
    """
    Get file metadata
    """
    file_stats = os.stat(file_path)
    file_ext = os.path.splitext(file_path)[1].lower()
    
    return {
        "filename": os.path.basename(file_path),
        "extension": file_ext,
        "size_bytes": file_stats.st_size,
        "size_mb": round(file_stats.st_size / (1024 * 1024), 2),
        "created_at": file_stats.st_ctime,
        "modified_at": file_stats.st_mtime
    }

def sanitize_filename(filename: str) -> str:
    """
    Sanitize filename to prevent security issues
    """
    # Remove path components
    filename = os.path.basename(filename)
    
    # Remove special characters except dots, hyphens, and underscores
    filename = re.sub(r'[^\w\s.-]', '', filename)
    
    # Replace spaces with underscores
    filename = filename.replace(' ', '_')
    
    # Limit length
    name, ext = os.path.splitext(filename)
    if len(name) > 100:
        name = name[:100]
    
    return name + ext

def chunk_text_for_processing(text: str, chunk_size: int = 1000, overlap: int = 100) -> list:
    """
    Split text into overlapping chunks for better context preservation
    """
    words = text.split()
    chunks = []
    
    for i in range(0, len(words), chunk_size - overlap):
        chunk = ' '.join(words[i:i + chunk_size])
        chunks.append(chunk)
    
    return chunks

def count_words(text: str) -> int:
    """Count words in text"""
    return len(text.split())

def estimate_reading_time(text: str, words_per_minute: int = 200) -> int:
    """
    Estimate reading time in minutes
    """
    word_count = count_words(text)
    return max(1, round(word_count / words_per_minute))

def preview_text(text: str, max_length: int = 200) -> str:
    """
    Get a preview of the text (first N characters)
    """
    if len(text) <= max_length:
        return text
    
    return text[:max_length] + "..."

def validate_text_content(text: str, min_words: int = 10) -> bool:
    """
    Validate that text has minimum content for processing
    """
    word_count = count_words(text)
    return word_count >= min_words and len(text.strip()) > 0

def format_error_message(error: Exception, context: str = "") -> str:
    """
    Format error messages for user-friendly display
    """
    error_msg = str(error)
    
    if "No such file" in error_msg:
        return f"File not found. Please check the file path."
    elif "Permission denied" in error_msg:
        return f"Permission denied. Cannot access the file."
    elif "not a valid" in error_msg.lower():
        return f"Invalid file format. Please upload a valid TXT, PDF, or DOCX file."
    elif "extract" in error_msg.lower():
        return f"Could not extract text from the file. The file might be corrupted or empty."
    else:
        return f"{context}: {error_msg}" if context else error_msg